import cv2
import paddlehub as hub
from docx import Document
import os
from threading import Thread  # 创建线程的模块
import pyttsx3 
from docx2pdf import convert

def say_task(content):
    engine = pyttsx3.init()
    engine.say(content) 
    engine.runAndWait()

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None

def delete_docx_prefix_description(docx):
    # delete_paragraph(docx.tables[0]) # 删除word中第一个table
    for p in docx.paragraphs:
        delete_paragraph(p)
        # if ''.join(p.text.split(' ')).lower()=='header_keyword':
        #     break
    # for p in docx.paragraphs:  
    #     if p.text.lower()=='': # 删除word中在开始部分的空白段落
    #         delete_paragraph(p)
    #     else:
    #         break

# 输出文件格式
    # print('准备开始语音播报...')
# 设置要播报的Unicode字符串
    # if(engine._inLoop):
    #     print('说的话冲突啦')
    #     say_list.append('刚刚我在说别的东西，我刚要说的东西是'+content)
    # else:


def predict(img,format_file):
    # 选择chinese_ocr_db_crnn_mobile模型
    tips=''
    if(format_file==0):#可能会和姿势识别的语音冲突
        raise RuntimeError("还没有决定格式！默认为word把")
        tips+="还没有决定格式,默认为word把"
        format_file=1
    if(format_file==1):
        tips+='保存格式为word'
    else:
        tips+='保存格式为pdf'
    say_task(tips)
    ocr = hub.Module(name="chinese_ocr_db_crnn_mobile")
    np_images = [img]
    results = ocr.recognize_text(
        images=np_images,  # 图片数据，nparray.shape 为 [H, W, C]，BGR格式；
        use_gpu=False,  # 是否使用 GPU；若使用GPU，请先设置CUDA_VISIBLE_DEVICES环境变量
        output_dir='ocr_result',  # 图片的保存路径
        visualization=False,  # 是否将识别结果保存为图片文件；
        box_thresh=0.5,  # 检测文本框置信度的阈值；
        text_thresh=0.5)  # 识别中文文本置信度的阈值；

    document = Document('test.docx')
    delete_docx_prefix_description(document)
    for result in results:
        data = result['data']
        for infomation in data:
            #保存到docx文件中（无排版）
            # print(os.path.dirname(__file__)+'\test.docx')
            document.add_paragraph(infomation['text'])
            document.save('test.docx')
    if (format_file==1):
        return
    elif (format_file==2):
        file = open('test.pdf', "w")
        file.close()
        convert('test.docx', 'test.pdf')
    else :
        raise Exception('格式未知错误？跑飞了')


# def predict(img="test.jpg"):
#     # 选择chinese_ocr_db_crnn_mobile模型
#     ocr = hub.Module(name="chinese_ocr_db_crnn_mobile")
#     np_images = [cv2.imread(img)]
#     results = ocr.recognize_text(
#         images=np_images,  # 图片数据，nparray.shape 为 [H, W, C]，BGR格式；
#         use_gpu=False,  # 是否使用 GPU；若使用GPU，请先设置CUDA_VISIBLE_DEVICES环境变量
#         output_dir='ocr_result',  # 图片的保存路径
#         visualization=False,  # 是否将识别结果保存为图片文件；
#         box_thresh=0.5,  # 检测文本框置信度的阈值；
#         text_thresh=0.5)  # 识别中文文本置信度的阈值；

#     for result in results:
#         data = result['data']
#         for infomation in data:
#             #保存到docx文件中（无排版）
#             # print(os.path.dirname(__file__)+'\test.docx')
#             document = Document('test.docx')
#             document.add_paragraph(infomation['text'])
#             document.save('test.docx')

if __name__ == '__main__':
    predict()
